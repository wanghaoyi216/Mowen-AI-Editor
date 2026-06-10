"""章节导出服务：Markdown / DOCX / PDF / TXT 四种格式。

设计要点：
  * 优先使用 ``chapter.final_content``，缺失时回退 ``draft_content``。
  * 章节元信息（编号 / 标题 / 摘要 / objective / 冲突 / 角色 / scenes）作为 front matter
    或 header 注入到各格式中。
  * DOCX 使用 ``python-docx``，PDF 使用 ``reportlab``。
    两个库都未安装时降级返回 ``text/plain`` 并记录 warning。
"""

from __future__ import annotations

import io
import logging
import re
from typing import Iterable, List

from app.models.chapter import Chapter
from app.models.plot_line import PlotLine
from app.schemas.chapter import ChapterRead


logger = logging.getLogger(__name__)


_CHAPTER_TAG = "chapter_scene_for:"


def _chapter_scenes(db, chapter: Chapter) -> List[PlotLine]:
    """拉取该章节下的所有 scene（按 priority 升序）。"""
    all_plots = list(
        db.query(PlotLine)
        .filter(PlotLine.project_id == chapter.project_id, PlotLine.plot_type == "chapter_scene")
        .order_by(PlotLine.priority.asc())
        .all()
    )
    matched = [p for p in all_plots if p.goal and f"{_CHAPTER_TAG}{chapter.id}" in p.goal]
    matched.sort(key=lambda p: (p.priority or 0, p.id or 0))
    return matched


def _safe(value, default: str = "") -> str:
    if value is None:
        return default
    return str(value)


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------
def render_markdown(chapter: Chapter, scenes: list[PlotLine]) -> str:
    content = (chapter.final_content or chapter.draft_content or "").strip()
    frontmatter = (
        "---\n"
        f"chapter_no: {chapter.chapter_no}\n"
        f'title: "{_safe(chapter.title)}"\n'
        f"word_count: {chapter.word_count or 0}\n"
        f"version: {chapter.version or 1}\n"
        f'objective: "{_safe(chapter.objective)[:200]}"\n'
        f'conflict: "{_safe(chapter.conflict)[:200]}"\n'
        "---\n\n"
    )
    body = f"# {_safe(chapter.title) or f'第{chapter.chapter_no}章'}\n\n"
    if chapter.objective:
        body += f"> **本章目标**：{chapter.objective}\n\n"
    if chapter.conflict:
        body += f"> **核心冲突**：{chapter.conflict}\n\n"
    if scenes:
        body += "## 本章场景\n\n"
        for i, sc in enumerate(scenes, 1):
            body += f"### 场景 {i}：{sc.title}\n\n"
            if sc.summary:
                body += f"{sc.summary}\n\n"
            if sc.goal:
                body += f"- 目标：{sc.goal}\n"
            if sc.conflict:
                body += f"- 冲突：{sc.conflict}\n"
            if sc.stakes:
                body += f"- 角色：{sc.stakes}\n"
            body += "\n"
    body += "---\n\n" + content + "\n"
    return frontmatter + body


# ---------------------------------------------------------------------------
# Plain text
# ---------------------------------------------------------------------------
def render_text(chapter: Chapter, scenes: list[PlotLine]) -> str:
    content = (chapter.final_content or chapter.draft_content or "").strip()
    lines: list[str] = [
        f"第 {chapter.chapter_no} 章  {_safe(chapter.title) or '未命名'}",
        "=" * 60,
        "",
    ]
    if chapter.objective:
        lines.append(f"【本章目标】{chapter.objective}")
        lines.append("")
    if chapter.conflict:
        lines.append(f"【核心冲突】{chapter.conflict}")
        lines.append("")
    if scenes:
        lines.append("【本章场景】")
        for i, sc in enumerate(scenes, 1):
            lines.append(f"  场景 {i}：{sc.title}")
            if sc.summary:
                lines.append(f"    {sc.summary}")
        lines.append("")
    lines.append("-" * 60)
    lines.append(content)
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
def render_docx(chapter: Chapter, scenes: list[PlotLine]) -> bytes:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt
    except ImportError:
        logger.warning("python-docx not installed, fallback to plain text")
        return render_text(chapter, scenes).encode("utf-8")

    doc = Document()
    # 标题
    title = _safe(chapter.title) or f"第{chapter.chapter_no}章"
    heading = doc.add_heading(f"第{chapter.chapter_no}章  {title}", level=0)
    heading.alignment = 1  # center
    # 元信息
    info = doc.add_paragraph()
    info.add_run(f"字数：{chapter.word_count or 0}    版本：{chapter.version or 1}\n").font.size = Pt(10)
    if chapter.objective:
        p = doc.add_paragraph()
        p.add_run("本章目标：").bold = True
        p.add_run(chapter.objective)
    if chapter.conflict:
        p = doc.add_paragraph()
        p.add_run("核心冲突：").bold = True
        p.add_run(chapter.conflict)
    # 场景列表
    if scenes:
        doc.add_heading("本章场景", level=1)
        for i, sc in enumerate(scenes, 1):
            doc.add_heading(f"场景 {i}：{sc.title}", level=2)
            if sc.summary:
                doc.add_paragraph(sc.summary)
            if sc.goal:
                p = doc.add_paragraph()
                p.add_run("目标：").bold = True
                p.add_run(sc.goal)
            if sc.conflict:
                p = doc.add_paragraph()
                p.add_run("冲突：").bold = True
                p.add_run(sc.conflict)
            if sc.stakes:
                p = doc.add_paragraph()
                p.add_run("角色：").bold = True
                p.add_run(sc.stakes)
    # 正文
    doc.add_heading("正文", level=1)
    content = (chapter.final_content or chapter.draft_content or "").strip()
    for para in re.split(r"\n\n+", content):
        if para.strip():
            doc.add_paragraph(para.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------
def render_pdf(chapter: Chapter, scenes: list[PlotLine]) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, PageBreak
        )
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    except ImportError:
        logger.warning("reportlab not installed, fallback to plain text")
        return render_text(chapter, scenes).encode("utf-8")

    # 注册中文字体（reportlab 自带 STSong-Light）
    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        chinese_font = "STSong-Light"
    except Exception:
        chinese_font = "Helvetica"

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleCN", parent=styles["Title"],
        fontName=chinese_font, fontSize=20, leading=26, alignment=1,
    )
    h1_style = ParagraphStyle(
        "H1CN", parent=styles["Heading1"],
        fontName=chinese_font, fontSize=16, leading=22, spaceBefore=10, spaceAfter=6,
    )
    h2_style = ParagraphStyle(
        "H2CN", parent=styles["Heading2"],
        fontName=chinese_font, fontSize=14, leading=20, spaceBefore=8, spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "BodyCN", parent=styles["BodyText"],
        fontName=chinese_font, fontSize=11, leading=18, spaceAfter=6,
    )
    meta_style = ParagraphStyle(
        "MetaCN", parent=styles["BodyText"],
        fontName=chinese_font, fontSize=10, leading=15,
        textColor="#666666",
    )

    story: list = []
    title = _safe(chapter.title) or f"第{chapter.chapter_no}章"
    story.append(Paragraph(f"第{chapter.chapter_no}章  {title}", title_style))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(
        f"字数：{chapter.word_count or 0}    版本：{chapter.version or 1}",
        meta_style,
    ))

    if chapter.objective:
        story.append(Paragraph("本章目标", h1_style))
        story.append(Paragraph(chapter.objective, body_style))
    if chapter.conflict:
        story.append(Paragraph("核心冲突", h1_style))
        story.append(Paragraph(chapter.conflict, body_style))

    if scenes:
        story.append(Paragraph("本章场景", h1_style))
        for i, sc in enumerate(scenes, 1):
            story.append(Paragraph(f"场景 {i}：{sc.title}", h2_style))
            if sc.summary:
                story.append(Paragraph(sc.summary, body_style))
            if sc.goal:
                story.append(Paragraph(f"<b>目标：</b>{sc.goal}", body_style))
            if sc.conflict:
                story.append(Paragraph(f"<b>冲突：</b>{sc.conflict}", body_style))
            if sc.stakes:
                story.append(Paragraph(f"<b>角色：</b>{sc.stakes}", body_style))

    story.append(Paragraph("正文", h1_style))
    content = (chapter.final_content or chapter.draft_content or "").strip()
    for para in re.split(r"\n\n+", content):
        if para.strip():
            story.append(Paragraph(para.strip().replace("\n", "<br/>"), body_style))

    doc.build(story)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------
SUPPORTED_FORMATS = ("md", "docx", "pdf", "txt")


def render_chapter(
    db,
    chapter: Chapter,
    format: str = "md",
) -> tuple[bytes, str, str]:
    """返回 ``(bytes, media_type, filename)`` 三元组。"""
    fmt = (format or "md").lower()
    if fmt not in SUPPORTED_FORMATS:
        raise ValueError(f"unsupported export format: {format}")

    scenes = _chapter_scenes(db, chapter)
    safe_title = _safe(chapter.title) or f"chapter_{chapter.chapter_no}"
    # 清理文件名特殊字符
    safe_title = "".join("_" if c in '<>:"/\\|?*' else c for c in safe_title).strip() or f"chapter_{chapter.chapter_no}"

    if fmt == "md":
        data = render_markdown(chapter, scenes).encode("utf-8")
        media = "text/markdown; charset=utf-8"
        ext = "md"
    elif fmt == "docx":
        data = render_docx(chapter, scenes)
        media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = "docx"
    elif fmt == "pdf":
        data = render_pdf(chapter, scenes)
        media = "application/pdf"
        ext = "pdf"
    else:  # txt
        data = render_text(chapter, scenes).encode("utf-8")
        media = "text/plain; charset=utf-8"
        ext = "txt"

    filename = f"chapter_{chapter.chapter_no:02d}_{safe_title}.{ext}"
    return data, media, filename
